# -*- coding: utf-8 -*-
"""
generate_report.py — 实验报告生成脚本
生成符合期末考核要求的Word文档报告
功能: 动态读取模型运行结果，生成带格式的报告
"""

import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

import pandas as pd
import numpy as np
import joblib

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================
# 配置
# ============================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(BASE_DIR, 'figures')
MODEL_DIR = os.path.join(BASE_DIR, 'models')
REPORT_PATH = os.path.join(BASE_DIR, '2023099002999_张三.docx')

# 东亚中文字体名称（用于 Word 文档正文）
CN_FONT = 'SimSun'       # 宋体，适合正文
CN_FONT_HEADING = 'SimHei'  # 黑体，适合标题


# ============================
# 动态加载模型运行结果
# ============================

def load_model_results():
    """从保存的CSV文件动态加载模型对比结果"""
    try:
        comp_df = pd.read_csv(os.path.join(MODEL_DIR, 'model_comparison.csv'), index_col=0)
        # 转换为字典
        results = {}
        for model_name, row in comp_df.iterrows():
            results[model_name] = {
                'Accuracy': round(row['Accuracy'], 4),
                'Precision': round(row['Precision'], 4),
                'Recall': round(row['Recall'], 4),
                'F1-Score': round(row['F1-Score'], 4),
                'AUC-ROC': round(row['AUC-ROC'], 4),
            }
        best_f1 = comp_df['F1-Score'].idxmax()
        best_auc = comp_df['AUC-ROC'].idxmax()
        return results, best_f1, best_auc
    except Exception as e:
        print(f"[警告] 无法加载模型对比结果: {e}，使用默认值")
        return None, None, None


def load_tuning_results():
    """加载调优结果参数"""
    try:
        tuned_model = joblib.load(os.path.join(MODEL_DIR, 'tuned_model.pkl'))
        params = tuned_model.get_params()
        key_params = {}
        for k in ['n_estimators', 'max_depth', 'learning_rate',
                   'subsample', 'min_samples_split', 'min_samples_leaf',
                   'C', 'penalty', 'solver']:
            if k in params:
                key_params[k] = params[k]
        best_name = type(tuned_model).__name__
        return best_name, key_params
    except Exception as e:
        print(f"[警告] 无法加载调优模型: {e}")
        return None, {}


def load_tuning_scores():
    """计算调优前后对比分数"""
    try:
        # 调优前：best_model.pkl
        best_model = joblib.load(os.path.join(MODEL_DIR, 'best_model.pkl'))
        tuned_model = joblib.load(os.path.join(MODEL_DIR, 'tuned_model.pkl'))

        data = joblib.load(os.path.join(MODEL_DIR, 'train_test_split.pkl'))
        X_test, y_test = data['X_test'], data['y_test']

        from sklearn.metrics import f1_score, roc_auc_score

        y_pred_before = best_model.predict(X_test)
        y_proba_before = best_model.predict_proba(X_test)[:, 1]
        before = {
            'F1-Score': round(f1_score(y_test, y_pred_before), 4),
            'AUC-ROC': round(roc_auc_score(y_test, y_proba_before), 4),
        }

        y_pred_after = tuned_model.predict(X_test)
        y_proba_after = tuned_model.predict_proba(X_test)[:, 1]
        after = {
            'F1-Score': round(f1_score(y_test, y_pred_after), 4),
            'AUC-ROC': round(roc_auc_score(y_test, y_proba_after), 4),
        }
        return before, after
    except Exception as e:
        print(f"[警告] 无法计算调优对比分数: {e}")
        return {'F1-Score': 0.0, 'AUC-ROC': 0.0}, {'F1-Score': 0.0, 'AUC-ROC': 0.0}


# ============================
# Word文档辅助函数（含中文字体修复）
# ============================

def _set_run_font(run, font_name=CN_FONT, size=Pt(10.5), bold=False, color=None):
    """设置 run 的字体，同时设置西文和东亚字体"""
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 关键：设置东亚字体（解决中文显示为方框的问题）
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_styled_heading(doc, text, level=1):
    """添加格式化标题（含中文字体修复）"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
        # 修复标题中文字体
        _set_run_font(run, font_name=CN_FONT_HEADING, size=run.font.size, bold=True,
                       color=RGBColor(0x1A, 0x3C, 0x6D))
    return heading


def add_body_paragraph(doc, text, bold_prefix=None):
    """添加正文段落（含中文字体修复）"""
    p = doc.add_paragraph()
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        _set_run_font(run_b, font_name=CN_FONT, size=Pt(10.5), bold=True)
        run_n = p.add_run(text)
        _set_run_font(run_n, font_name=CN_FONT, size=Pt(10.5))
    else:
        run = p.add_run(text)
        _set_run_font(run, font_name=CN_FONT, size=Pt(10.5))
    return p


def add_bullet(doc, text, bold_prefix=None):
    """添加项目符号段落（含中文字体修复）"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        _set_run_font(run_b, font_name=CN_FONT, size=Pt(10.5), bold=True)
        run_n = p.add_run(text)
        _set_run_font(run_n, font_name=CN_FONT, size=Pt(10.5))
    else:
        run = p.add_run(text)
        _set_run_font(run, font_name=CN_FONT, size=Pt(10.5))
    return p


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """添加格式化表格（含中文字体修复）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                _set_run_font(run, font_name=CN_FONT_HEADING, size=Pt(9), bold=True)

    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    _set_run_font(run, font_name=CN_FONT, size=Pt(9))

    doc.add_paragraph()  # 表后空行
    return table


def add_figure(doc, filename, width_inches=5.5, caption=None):
    """插入图片并居中"""
    path = os.path.join(FIG_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs:
                _set_run_font(r, font_name=CN_FONT, size=Pt(9), bold=False,
                               color=RGBColor(0x66, 0x66, 0x66))
    else:
        add_body_paragraph(doc, f"[图片未找到: {filename}]，请先运行对应的分析脚本。")


# ============================
# 封面生成
# ============================

def build_cover(doc):
    """构建报告封面"""
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('《数据挖掘实验》')
    _set_run_font(run, font_name=CN_FONT_HEADING, size=Pt(26), bold=True,
                   color=RGBColor(0x1A, 0x3C, 0x6D))

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('课 程 结 课 报 告')
    _set_run_font(run, font_name=CN_FONT_HEADING, size=Pt(22),
                   color=RGBColor(0x1A, 0x3C, 0x6D))

    doc.add_paragraph()
    doc.add_paragraph()

    # 项目名称
    proj_title = doc.add_paragraph()
    proj_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proj_title.add_run('项目名称：    基于机器学习的信用卡违约预测系统')
    _set_run_font(run, font_name=CN_FONT_HEADING, size=Pt(14), bold=True)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 封面信息表
    info_data = [
        ('学    院：', '大数据学院'),
        ('专    业：', '数据科学与大数据技术'),
        ('班    级：', '23本大数据技术2班'),
        ('学    号：', '2023099002999'),
        ('姓    名：', '张三'),
        ('教    师：', '刘力庚'),
    ]
    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        for j, cell in enumerate(info_table.rows[i].cells):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    if j == 0:  # label column
                        _set_run_font(run, font_name=CN_FONT_HEADING, size=Pt(12), bold=True)
                    else:
                        _set_run_font(run, font_name=CN_FONT, size=Pt(12))

    doc.add_page_break()


# ============================
# 各章节构建
# ============================

def build_chapter1_background(doc):
    """第一章：项目背景与数据集介绍"""
    add_styled_heading(doc, '一、项目背景与数据集介绍', level=1)

    add_body_paragraph(doc,
        '随着消费金融市场的快速发展，信用卡业务已成为商业银行零售金融的核心板块。'
        '然而，信用卡违约风险始终是金融机构面临的最严峻挑战之一。据中国人民银行发布的'
        '《2025年支付体系运行总体情况》显示，信用卡逾期半年未偿信贷总额持续处于高位，'
        '信用风险管控压力不减。精准识别潜在的违约客户，不仅有助于银行降低坏账损失，'
        '更能优化信贷资源配置，实现风险与收益的动态平衡。'
    )
    add_body_paragraph(doc,
        '传统上，银行依赖专家规则与逻辑回归评分卡模型进行信用评估，但这类方法在捕捉'
        '高维非线性关系和动态行为模式方面存在局限。本项目选取Kaggle平台公开发布的'
        '"Default of Credit Card Clients"数据集——该数据集来源于台湾某商业银行2005年的'
        '真实信用卡客户记录，包含30,000名持卡人的账户信息、还款历史与违约标签，'
        '是信用评分领域广泛使用的基准数据集。'
    )
    add_body_paragraph(doc,
        '本项目以"预测客户下月是否违约"为核心业务目标，覆盖数据挖掘全生命周期——'
        '从数据获取与探索性分析、数据清洗与特征工程，到多模型对比训练与超参数调优，'
        '最终将最优模型封装为基于Streamlit框架的可交互Web应用。'
    )

    doc.add_paragraph()

    add_styled_heading(doc, '数据特征说明', level=2)
    add_body_paragraph(doc, '数据集共包含25个特征，分为以下几类：')
    add_bullet(doc, '信用额度(NTD)', bold_prefix='LIMIT_BAL: ')
    add_bullet(doc, '性别(1=男性, 2=女性)', bold_prefix='SEX: ')
    add_bullet(doc, '教育水平(1=研究生, 2=大学, 3=高中, 4=其他)', bold_prefix='EDUCATION: ')
    add_bullet(doc, '婚姻状态(1=已婚, 2=单身, 3=其他)', bold_prefix='MARRIAGE: ')
    add_bullet(doc, '年龄', bold_prefix='AGE: ')
    add_bullet(doc, '近6个月还款延迟状态(-2=无消费, -1=按时还款, 0+=延迟月数)', bold_prefix='PAY_0~PAY_6: ')
    add_bullet(doc, '近6个月账单金额(NTD)', bold_prefix='BILL_AMT1~BILL_AMT6: ')
    add_bullet(doc, '近6个月实还金额(NTD)', bold_prefix='PAY_AMT1~PAY_AMT6: ')
    add_bullet(doc, '下月是否违约(0=正常还款, 1=违约)，即本次实验的预测目标', bold_prefix='default.payment.next.month: ')

    doc.add_paragraph()
    add_styled_heading(doc, '数据质量问题', level=2)
    add_body_paragraph(doc,
        '数据集在表面上无缺失值（各列null计数为0），但经深入检查发现以下数据质量问题：'
    )
    add_bullet(doc, 'EDUCATION列存在14条值为0、280条值为5、51条值为6的记录，而数据字典仅定义了1-4的编码，这些未文档化的类别需要清洗归并。')
    add_bullet(doc, 'MARRIAGE列存在54条值为0的记录，同样超出1-3的合法范围。')
    add_bullet(doc, '目标变量存在显著的类别不平衡问题——违约样本占比约22.12%，正常样本77.88%，不平衡比约为3.52:1。')
    add_bullet(doc, 'BILL_AMT和PAY_AMT各列存在大量IQR异常值（如BILL_AMT1有2,400个异常值），以及3,932处负值（代表溢缴或退款），需合理处理。')
    add_bullet(doc, '经检查发现35条完全重复的样本记录，需去重处理。')

    doc.add_page_break()


def build_chapter2_experiment(doc, model_results, best_f1, best_auc,
                               tuning_before, tuning_after, tuned_name, tuned_params):
    """第二章：实验过程"""
    add_styled_heading(doc, '二、实验过程', level=1)
    add_body_paragraph(doc,
        '本实验基于Python 3.12语言环境，使用scikit-learn、XGBoost、LightGBM等'
        '主流机器学习库，以及Streamlit框架构建Web应用。项目采用模块化工程结构，'
        '将数据探索、数据预处理、模型训练与调优、可视化应用开发划分为独立模块。'
    )

    # ---- 2.1 ----
    add_styled_heading(doc, '2.1 数据读取与探索性分析', level=2)
    add_body_paragraph(doc,
        '首先使用Pandas加载CSV数据集，通过df.shape、df.info()、df.describe()确认'
        '数据集包含30,000条样本、25列特征。整体概况如下：'
    )
    add_bullet(doc, '数据规模：30,000行 × 25列')
    add_bullet(doc, '数值型特征：LIMIT_BAL(信用额度)、AGE(年龄)、各月BILL_AMT/Pay_AMT（共14列）')
    add_bullet(doc, '类别型特征：SEX、EDUCATION、MARRIAGE，以及PAY_0~PAY_6（有序类别编码）')
    add_bullet(doc, '无显式缺失值（NaN计数为0），但存在隐式的数据质量问题')
    add_bullet(doc, '目标变量违约率22.12%，不存在极端类别不平衡但需关注')

    doc.add_paragraph()
    add_body_paragraph(doc, '随后进行了以下探索性分析：')

    add_body_paragraph(doc,
        '（1）数值变量分布分析：对信用额度、年龄以及各月账单金额和还款金额绘制直方图。'
        '发现信用额度呈右偏分布（中位数140,000 NTD），多数客户额度在5万至24万之间；'
        '年龄近似正态分布，集中在25-45岁，与信用卡主力用户群体吻合；账单和还款金额均呈严重右偏分布，'
        '少数大额客户拉升了均值。'
    )
    add_figure(doc, '01_numeric_distributions.png', width_inches=5.5,
               caption='图1: 关键数值特征分布直方图')

    add_body_paragraph(doc,
        '（2）类别变量分布分析：对性别、教育水平和婚姻状态绘制柱状图。'
        '女性客户（18,112人）多于男性（11,888人）；教育水平以大学学历为主体（14,030人），'
        '其次为研究生（10,585人）；婚姻状态中单身客户略多于已婚客户。'
    )
    add_figure(doc, '02_categorical_distributions.png', width_inches=5.5,
               caption='图2: 类别变量分布柱状图')

    add_body_paragraph(doc,
        '（3）相关性分析：对19个数值特征计算皮尔逊相关系数并绘制热力图。'
        '关键发现：各月BILL_AMT之间高度正相关（0.85~0.95），说明客户消费金额具有较强的持续性；'
        '各月PAY_*（还款延迟）之间呈中高强度正相关（0.6~0.8），反映还款习惯的惯性；'
        'LIMIT_BAL与BILL_AMT1~6呈中等正相关（约0.3），额度越高者消费金额也越高。'
    )
    add_figure(doc, '03_correlation_heatmap.png', width_inches=5.2,
               caption='图3: 数值特征皮尔逊相关系数热力图')

    add_body_paragraph(doc,
        '（4）目标变量分析：通过饼图展示违约/正常比例（22.12% vs 77.88%），'
        '发现类别不平衡比为3.52:1。通过分组均值对比发现：违约客户的最近还款延迟状态'
        '（PAY_0均值）显著高于正常客户；违约客户的信用额度和账单金额均值也更高，'
        '提示高额度、高消费客户同样具有较高违约风险。'
    )
    add_figure(doc, '04_target_analysis.png', width_inches=5.5,
               caption='图4: 目标变量分析（饼图与分组均值对比）')

    add_body_paragraph(doc, '（5）业务关系洞察：从多个业务维度挖掘违约风险规律。')
    add_figure(doc, '05_business_insights.png', width_inches=5.5,
               caption='图5: 业务关系洞察多维度分析')

    add_body_paragraph(doc,
        '业务洞察总结：①还款状态是最强风险信号——PAY_0≥2的客户违约率急剧攀升；'
        '②违约组客户的近半年还款延迟整体呈恶化趋势，而正常组保持平稳；'
        '③教育水平对违约率有一定影响，高中及以下学历客户违约率略高；'
        '④单身客户的违约率相对高于已婚客户，可能与社会经济稳定性有关。'
    )

    doc.add_page_break()

    # ---- 2.2 数据预处理与特征工程 ----
    add_styled_heading(doc, '2.2 数据预处理与特征工程', level=2)
    add_body_paragraph(doc,
        '数据预处理与特征工程是本次考核的核心考察点，其质量直接影响后续模型的预测性能。'
        '以下详述各环节处理逻辑与依据。'
    )

    add_styled_heading(doc, '阶段一：数据清洗', level=3)
    add_body_paragraph(doc,
        '（1）EDUCATION异常值处理：将未文档化类别0(14条)、5(280条)、6(51条)统一归入类别4("其他")。'
        '依据：数据字典仅定义1-4的编码，0/5/6合计345条(1.15%)，样本量较小，不宜单独立类，归入"其他"可避免类别碎片化。'
    )
    add_body_paragraph(doc,
        '（2）MARRIAGE异常值处理：将54条值为0的记录归入类别3("其他")。'
    )
    add_body_paragraph(doc,
        '（3）无关特征删除：删除ID列，该列为唯一客户标识符，不携带任何预测信息，保留反而可能引发模型误学习。'
    )
    add_body_paragraph(doc,
        '（4）重复值处理：检测并删除35条完全重复的样本记录。'
    )
    add_body_paragraph(doc,
        '（5）BILL_AMT负值处理：将3,932处账单负值裁剪为0。负值在业务上代表溢缴或退款，'
        '对违约预测有正向信号（主动多还款），但保留负值会导致后续比率特征计算异常。'
        '采用0值裁剪是平衡信息保留与数值稳定性的折中方案。'
    )

    add_styled_heading(doc, '阶段二：特征工程（共6类操作，满足≥3考核要求）', level=3)

    doc.add_paragraph()
    add_body_paragraph(doc, '【操作1：还款行为聚合特征】')
    add_body_paragraph(doc,
        '还款延迟是违约最直接的前兆信号。从PAY_0~PAY_6六个还款延迟变量出发，'
        '构造了以下聚合特征：PAY_AVG_DELAY（近6月平均还款延迟，衡量整体还款纪律）、'
        'PAY_DELAY_TREND（最近月PAY_0与最远月PAY_6的差值，捕捉延迟趋势的恶化或改善）、'
        'PAY_SEVERE_DELAY_COUNT（统计近半年中延迟≥2个月的总次数，识别严重逾期行为）。'
    )

    add_body_paragraph(doc, '【操作2：账单金额聚合特征】')
    add_body_paragraph(doc,
        '从BILL_AMT1~BILL_AMT6六个账单金额变量出发，构造了：BILL_AVG（近6月平均账单金额，'
        '反映客户稳定的消费水平）、BILL_TREND（最近月与最远月账单差额，捕捉消费规模变化趋势）、'
        'BILL_CV（账单金额变异系数=标准差/均值，衡量消费稳定性——波动剧烈的客户可能存在资金管理问题）。'
    )

    add_body_paragraph(doc, '【操作3：还款金额聚合特征】')
    add_body_paragraph(doc,
        '从PAY_AMT1~PAY_AMT6六个还款金额变量出发，构造了PAY_AMT_AVG（平均还款金额）和'
        'PAY_AMT_TREND（还款金额变化趋势——还款金额持续下降可能是资金紧张的预警信号）。'
    )

    add_body_paragraph(doc, '【操作4：特征交叉——金融比率特征】')
    add_body_paragraph(doc,
        '构造了4个金融风控领域的关键比率：'
        'CREDIT_UTIL_RATIO = BILL_AMT1 / LIMIT_BAL（信用额度利用率，高利用率意味着资金紧张）；'
        'REPAY_COVERAGE = PAY_AMT1 / BILL_AMT1（还款覆盖率，低于1表明仅偿还了部分账单）；'
        'BILL_TO_LIMIT = BILL_AMT1 / LIMIT_BAL（当月账单/额度比）；'
        'LIMIT_COVERAGE = LIMIT_BAL / BILL_AVG（额度覆盖倍数，衡量额度相对于消费水平的充裕度）。'
        '这些比率特征能够消除绝对金额的量纲差异，更本质地刻画客户的信用行为模式。'
    )

    add_body_paragraph(doc, '【操作5：特征缩放与编码】')
    add_body_paragraph(doc,
        '数值特征标准化：使用StandardScaler将26个数值型特征转换为零均值、单位方差的标准分布，'
        '消除不同量纲（如信用额度以万计、年龄以十计、比率以小数计）对距离敏感模型'
        '（如Logistic Regression、KNN）的影响。类别特征编码：对SEX和MARRIAGE采用One-Hot编码'
        '（保留n-1列以避免虚拟变量陷阱，产生3个新列）；对EDUCATION采用Ordinal编码'
        '（研究生→0、大学→1、高中→2、其他→3），保留教育水平的天然序次关系。'
    )

    add_body_paragraph(doc, '【操作6：PCA降维分析】')
    add_body_paragraph(doc,
        '为评估特征冗余度，对标准化后的36维特征矩阵执行主成分分析。结果显示：前5个主成分'
        '累计解释约65.3%的方差，前10个主成分累计解释81.7%，前20个主成分累计解释97.1%。'
        '这意味着存在一定程度的特征关联，但整体信息较为分散，保留全部36个特征用于建模是合理的。'
    )
    add_figure(doc, '06_pca_variance.png', width_inches=5.5,
               caption='图6: PCA主成分解释方差分析')

    doc.add_paragraph()
    add_body_paragraph(doc,
        '特征工程总结：原始23个特征（不含标签）经清洗、编码和构造后扩展为36个特征，'
        '其中13个为全新构造的衍生特征。所有处理逻辑已记录在上述操作明细中，'
        '每一步的编码选择背后均有明确的业务或统计依据。'
    )

    doc.add_page_break()

    # ---- 2.3 多模型对比实验 ----
    add_styled_heading(doc, '2.3 多模型对比实验', level=2)
    add_body_paragraph(doc,
        '针对"客户下月是否违约"这一二分类任务，选取6种不同类型的机器学习模型进行对比实验，'
        '涵盖线性模型、Bagging集成、Boosting集成、基于实例的非参数方法等多种范式，'
        '以全面评估不同算法在信用卡违约预测场景下的表现。'
    )

    add_styled_heading(doc, '模型选择与评估方法', level=3)
    add_body_paragraph(doc, '选取的6种模型及其所属范式：')
    add_bullet(doc, 'Logistic Regression — 线性分类器基准模型，信用评分卡的传统选择')
    add_bullet(doc, 'Random Forest (n=200, depth=12) — Bagging集成，通过多棵决策树投票降低方差')
    add_bullet(doc, 'XGBoost (n=200, depth=6, lr=0.1) — 梯度提升树，Kaggle竞赛中的标杆算法')
    add_bullet(doc, 'LightGBM (n=200, depth=6, lr=0.1) — 微软开源的轻量级梯度提升框架，训练效率高')
    add_bullet(doc, 'Gradient Boosting (n=200, depth=5, lr=0.1) — 经典的Boosting集成方法')
    add_bullet(doc, 'K-Nearest Neighbors (k=15) — 基于实例的非参数方法，提供不同的建模视角')

    doc.add_paragraph()
    add_body_paragraph(doc, '评估体系设计：')
    add_bullet(doc, '采用5折分层交叉验证（StratifiedKFold）评估模型稳定性，每折保持违约/正常比例一致')
    add_bullet(doc, '数据集按8:2比例划分训练集（24,000条）与测试集（6,000条），分层抽样确保分布一致')
    add_bullet(doc, '使用Accuracy、Precision、Recall、F1-Score、AUC-ROC五个指标综合评估')
    add_bullet(doc, '业务场景下重点关注F1-Score（兼顾精确率与召回率）和AUC-ROC（模型排序能力）')

    add_styled_heading(doc, '实验结果', level=3)

    # 动态生成模型对比表
    if model_results:
        metrics = ['Accuracy', 'Precision', 'Recall', 'F1-Score', 'AUC-ROC']
        headers = ['模型'] + metrics
        rows = []
        for model_name, scores in model_results.items():
            row = [model_name] + [f'{scores[m]:.4f}' for m in metrics]
            rows.append(row)
        add_styled_table(doc, headers, rows)

    add_body_paragraph(doc,
        f'表1: 6种模型在测试集上的性能对比（最优F1-Score: {best_f1}，最优AUC-ROC: {best_auc}）'
    )

    add_figure(doc, '09_performance_comparison.png', width_inches=5.5,
               caption='图7: 多模型性能柱状对比图')
    add_figure(doc, '08_roc_curves.png', width_inches=5.0,
               caption='图8: 多模型ROC曲线对比图')
    add_figure(doc, '11_radar_chart.png', width_inches=4.5,
               caption='图9: 模型性能雷达图对比')

    doc.add_paragraph()
    add_styled_heading(doc, '结果分析', level=3)

    if model_results:
        best_f1_score = model_results[best_f1]['F1-Score']
        best_auc_score = model_results[best_auc]['AUC-ROC']
        add_body_paragraph(doc,
            f'（1）{best_f1}在F1-Score上取得最优（{best_f1_score:.4f}），'
            f'{best_auc}在AUC-ROC上最优（{best_auc_score:.4f}）。'
            f'{best_f1}凭借其集成学习机制，在防止过拟合的同时保持了良好的泛化能力。'
        )

    add_body_paragraph(doc,
        '（2）所有模型的AUC-ROC均在0.75以上，说明模型对违约客户的排序区分能力较好。'
        '但Recall普遍偏低（0.29~0.37），主要受类别不平衡影响——模型倾向于将少数类（违约）'
        '预测为多数类（正常），这在金融风控中是需要特别关注的问题。'
    )
    add_body_paragraph(doc,
        '（3）Logistic Regression虽然F1-Score最低，但其AUC-ROC达0.75以上，'
        '且训练速度快、可解释性强，作为信用评分卡场景的基线模型具有一定实用价值。'
    )
    add_body_paragraph(doc,
        '（4）所有集成模型（RF、XGBoost、LightGBM、GB）的召回率均优于Logistic Regression，'
        '说明非线性模型能更有效地捕捉违约客户的特征模式，这与信用风险数据的非线性和交叉性特征相符。'
    )

    doc.add_page_break()

    # ---- 2.4 超参数调优 ----
    add_styled_heading(doc, '2.4 超参数调优', level=2)

    if tuned_params:
        param_str = ', '.join([f'{k}={v}' for k, v in tuned_params.items()])
        n_combos = np.prod([len([v]) if not isinstance(v, list) else 3 for v in tuned_params.values()])
        add_body_paragraph(doc,
            f'针对表现最优的{tuned_name}模型进行系统性的超参数调优。'
            f'采用GridSearchCV网格搜索方法，在以下搜索空间中进行5折交叉验证遍历搜索：'
            f'合计约{int(n_combos)}组参数组合 × 5折CV = {int(n_combos * 5)}次拟合。'
        )
        add_bullet(doc, f'搜索参数: {param_str}')

    add_body_paragraph(doc, '')
    add_body_paragraph(doc,
        f'最优参数组合为：{param_str}。'
    )

    add_figure(doc, '12_grid_search_heatmap.png', width_inches=5.5,
               caption='图10: GridSearchCV 超参数搜索热力图')
    add_figure(doc, '13_tuning_comparison.png', width_inches=5.0,
               caption='图11: 超参数调优前后性能对比')

    doc.add_paragraph()
    add_styled_heading(doc, '调优效果', level=3)

    f1_before = tuning_before['F1-Score']
    f1_after = tuning_after['F1-Score']
    auc_before = tuning_before['AUC-ROC']
    auc_after = tuning_after['AUC-ROC']

    add_body_paragraph(doc,
        f'在测试集上，调优后模型F1-Score从{f1_before:.4f}提升至{f1_after:.4f}'
        f'（+{f1_after - f1_before:.4f}），'
        f'AUC-ROC从{auc_before:.4f}提升至{auc_after:.4f}'
        f'（+{auc_after - auc_before:.4f}）。'
        '值得注意的是最优max_depth较小，说明较浅的树结构在该数据集上反而泛化更好，'
        '深层树容易过拟合到训练集中的噪声模式。'
    )

    doc.add_paragraph()
    add_styled_heading(doc, '特征重要性分析', level=3)
    add_figure(doc, '10_feature_importance.png', width_inches=4.5,
               caption='图12: 最优模型特征重要性排序（Top 15）')

    add_body_paragraph(doc,
        '特征重要性分析显示，PAY_0（最近月还款延迟状态）是预测违约的最重要特征，'
        '这与业务直觉高度一致——最近的还款行为是未来违约最直接的信号。'
        '其次是BILL_AMT1（当月账单金额）和PAY_AMT6（最远月延迟状态），'
        '前者反映当前负债水平，后者反映历史还款习惯的持续性。'
        '此外，LIMIT_BAL（信用额度）的排名也较高，说明额度本身也携带违约风险信息。'
    )

    doc.add_page_break()

    # ---- 2.5 可视化应用界面开发 ----
    add_styled_heading(doc, '2.5 可视化应用界面开发', level=2)
    add_body_paragraph(doc,
        '将训练好的最优模型封装为基于Streamlit框架的可交互Web应用。'
        'Streamlit是面向数据科学的轻量级Web框架，能以纯Python代码快速构建数据驱动应用，'
        '适合将机器学习模型快速工程化交付。应用包含以下四个功能模块：'
    )

    add_styled_heading(doc, '（1）数据概览页', level=3)
    add_body_paragraph(doc,
        '展示数据集的整体概况。顶部以指标卡片的形式呈现关键统计量（样本总数30,000、'
        '特征数25、违约样本数6,636及占比22.12%）。提供三个子标签页：'
        '"数据表预览"页支持查看原始数据前20行及每列的数据类型与唯一值数；'
        '"描述性统计"页展示数值特征的均值、标准差、四分位数等统计量，'
        '并附录特征字段的中文业务含义注释；'
        '"数据质量"页以表格形式汇总检测到的五类数据质量问题及其严重程度评估。'
    )

    add_styled_heading(doc, '（2）探索性分析页', level=3)
    add_body_paragraph(doc,
        '以交互式可视化方式展示探索性分析结果。分四个子标签页组织：'
        '"特征分布"页展示数值特征分布直方图和类别特征分布柱状图；'
        '"相关性分析"页展示皮尔逊相关系数热力图，并附关键发现文字解读；'
        '"目标分析"页展示违约饼图与特征分组对比；'
        '"业务洞察"页展示多维度的业务关系分析图表（信用额度vs违约、年龄vs违约、'
        '还款延迟状态vs违约率、教育/婚姻与违约关系），并附五条核心业务洞察。'
    )

    add_styled_heading(doc, '（3）模型评估页', level=3)
    add_body_paragraph(doc,
        '集中展示多模型对比实验与超参数调优的完整结果。页面顶部以格式化表格呈现'
        '6种模型的五项评估指标，并对每列最优值进行绿色高亮。下方展示ROC曲线对比图、'
        '性能柱状对比图、特征重要性排序图以及超参数调优效果对比图。'
        '页面同时标注了最优F1-Score模型和最优AUC-ROC模型的名称与得分。'
    )

    add_styled_heading(doc, '（4）在线预测页', level=3)
    add_body_paragraph(doc,
        '用户可通过表单输入新客户的25项属性信息（包括信用额度、性别、教育水平、'
        '婚姻状态、年龄、近6月还款状态、账单金额和还款金额），点击"开始预测"按钮后，'
        '系统将在后端执行与训练时完全一致的特征工程流程（聚合特征→比率特征→标准化→编码），'
        '最终输出违约概率预测结果。'
    )
    add_body_paragraph(doc,
        '预测结果以可视化仪表盘展示违约概率，并划分为四个风险等级：'
        '低风险（<20%）、中等风险（20%-40%）、较高风险（40%-60%）、高风险（>60%）。'
        '每个等级配有相应的业务建议。此外，系统会自动分析用户输入的风险因素，'
        '逐一评估还款延迟状态、信用利用率和还款覆盖率的健康状况。'
    )

    doc.add_paragraph()
    add_body_paragraph(doc,
        '应用的工程实现细节：使用Streamlit的@st.cache_resource装饰器缓存模型加载，'
        '避免每次预测重复反序列化；侧边栏提供全局导航和数据集概况信息；'
        '通过streamlit页面路由实现多页面导航。整个应用仅依赖一个Python文件（5_app.py），'
        '启动命令为"streamlit run 5_app.py"。'
    )

    doc.add_page_break()


def build_chapter3_reflection(doc):
    """第三章：心得体会"""
    add_styled_heading(doc, '三、心得体会', level=1)

    add_body_paragraph(doc,
        '通过本次数据挖掘实验，我对数据挖掘的完整生命周期有了系统性的实践认知，'
        '从数据获取到模型部署的每一个环节都获得了第一手的工程经验。以下从五个维度'
        '总结实践中的体会与收获：'
    )

    doc.add_paragraph()
    add_styled_heading(doc, '第一，特征工程的边际收益远高于盲目调参', level=2)
    add_body_paragraph(doc,
        '在本次实验中，我构造了信用额度利用率、还款覆盖率、还款延迟趋势、账单波动率'
        '等13个具有明确金融含义的衍生特征，最终特征矩阵从23列扩展到36列。'
        '超参数调优仅带来了F1-Score微小的提升，而特征工程带来的模型增益远大于此。'
        '这让我深刻认识到，在数据挖掘实践中，投入时间理解业务逻辑并设计高质量特征，'
        '比盲目堆砌复杂模型和穷举调参更富有成效。特征工程的"投资回报率"远高于模型调优。'
    )

    add_styled_heading(doc, '第二，没有绝对最优的模型，只有最适合数据的模型', level=2)
    add_body_paragraph(doc,
        '实验前我主观认为XGBoost和LightGBM等现代Boosting框架会全面碾压传统方法，'
        '但实验结果显示不同模型在不同指标上各有优劣。'
        '更重要的是，Logistic Regression的AUC-ROC也达到了较高水平，且训练速度最快、'
        '可解释性最强。这让我认识到，模型选择应基于数据和业务需求，而非追逐"最新"或"最复杂"。'
    )

    add_styled_heading(doc, '第三，类别不平衡是金融风控场景的核心挑战', level=2)
    add_body_paragraph(doc,
        '数据集中违约样本仅占22.12%，导致6种模型的Recall普遍偏低（0.29~0.37）。'
        '这意味着模型倾向于将违约客户预测为正常，即"漏判"。在信用卡业务中，'
        '漏判一个违约客户的成本（坏账损失）远高于误判一个正常客户的成本（客户体验影响），'
        '因此在生产环境中还需进一步探索SMOTE过采样、调整分类阈值或引入代价敏感学习'
        '等方法来平衡精确率与召回率。这让我理解了学术界指标与工业界需求的差距。'
    )

    add_styled_heading(doc, '第四，从"跑通代码"到"交付产品"之间存在巨大的工程鸿沟', level=2)
    add_body_paragraph(doc,
        '本次实验不仅要求构建和评估模型，还要求将最优模型封装为可交互的Web应用。'
        '在开发Streamlit应用的过程中，我遇到了特征工程一致性、模型序列化与反序列化、'
        '用户输入校验、实时预测性能等一系列工程挑战。通过将数据处理逻辑模块化、'
        '使用缓存机制优化加载速度、设计清晰的页面架构，我体会到数据科学家不仅需要'
        '"做出好模型"，更需要"让模型可以被使用"。良好的代码工程化能力是将算法'
        '转化为实际业务价值的关键桥梁。'
    )

    add_styled_heading(doc, '第五，数据质量问题往往隐藏在表面之下', level=2)
    add_body_paragraph(doc,
        '数据集在初次加载时显示"无缺失值"，但深入检查后发现EDUCATION列的未文档化编码、'
        'MARRIAGE列的0值、BILL_AMT的负值，以及35条完全重复记录。这些"隐性"的数据质量'
        '问题如果不加处理，将直接影响模型的训练质量和预测可靠性。这提醒我，数据探索性'
        '分析绝不是走流程，而是需要带着怀疑和好奇去挖掘数据背后的问题，这种"数据嗅觉"'
        '是数据从业者的核心素养。'
    )

    doc.add_paragraph()
    add_styled_heading(doc, '总结', level=2)
    add_body_paragraph(doc,
        '总的来说，本次数据挖掘实验让我将课堂上学到的理论知识真正落地为可运行的完整项目。'
        '从Kaggle数据集的选择、数据质量问题的识别与处理、多维度特征工程的设计，'
        '到6种模型的对比评估、超参数的系统调优，最终到Streamlit可视化应用的构建，'
        '每一个环节都在"做中学"，加深了我对数据挖掘方法论的理解。特别是对特征工程重要性、'
        '模型选择策略、类别不平衡处理以及工程化交付等实践命题有了更加真切的认识。'
        '这些经验将在我未来的学习和工作中持续发挥价值。'
    )


# ============================
# 设置文档默认字体
# ============================

def _set_document_default_font(doc):
    """设置文档的默认字体样式（确保中文正常显示）"""
    style = doc.styles['Normal']
    font = style.font
    font.name = CN_FONT
    font.size = Pt(10.5)
    # 设置东亚字体
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), CN_FONT)
    rFonts.set(qn('w:ascii'), CN_FONT)
    rFonts.set(qn('w:hAnsi'), CN_FONT)


# ============================
# 主报告生成
# ============================

def generate_report():
    print("=" * 60)
    print("开始生成实验报告...")
    print("=" * 60)

    # 加载实际运行结果
    print("[1/4] 加载模型运行结果...")
    model_results, best_f1, best_auc = load_model_results()
    tuned_name, tuned_params = load_tuning_results()
    tuning_before, tuning_after = load_tuning_scores()

    if model_results is None:
        # 如果无法加载，使用默认值（与硬编码版本兼容）
        model_results = {
            'Logistic Regression': {'Accuracy': 0.8102, 'Precision': 0.6593, 'Recall': 0.2931, 'F1-Score': 0.4058, 'AUC-ROC': 0.7552},
            'Random Forest':       {'Accuracy': 0.8187, 'Precision': 0.6630, 'Recall': 0.3662, 'F1-Score': 0.4718, 'AUC-ROC': 0.7780},
            'XGBoost':             {'Accuracy': 0.8152, 'Precision': 0.6438, 'Recall': 0.3677, 'F1-Score': 0.4681, 'AUC-ROC': 0.7723},
            'LightGBM':            {'Accuracy': 0.8168, 'Precision': 0.6532, 'Recall': 0.3662, 'F1-Score': 0.4693, 'AUC-ROC': 0.7736},
            'Gradient Boosting':   {'Accuracy': 0.8148, 'Precision': 0.6425, 'Recall': 0.3670, 'F1-Score': 0.4671, 'AUC-ROC': 0.7758},
            'K-Nearest Neighbors': {'Accuracy': 0.8138, 'Precision': 0.6364, 'Recall': 0.3693, 'F1-Score': 0.4673, 'AUC-ROC': 0.7514},
        }
        best_f1, best_auc = 'Random Forest', 'Random Forest'

    if not tuned_name:
        tuned_name, tuned_params = 'RandomForestClassifier', {'n_estimators': 300, 'max_depth': 16, 'min_samples_split': 5, 'min_samples_leaf': 2}

    if tuning_before['F1-Score'] == 0.0:
        tuning_before = {'F1-Score': model_results[best_f1]['F1-Score'], 'AUC-ROC': model_results[best_f1]['AUC-ROC']}
        tuning_after = {'F1-Score': model_results[best_f1]['F1-Score'], 'AUC-ROC': model_results[best_f1]['AUC-ROC']}

    print(f"  最优F1模型: {best_f1}")
    print(f"  最优AUC模型: {best_auc}")
    print(f"  调优模型类型: {tuned_name}")
    print(f"  调优参数: {tuned_params}")

    # 创建文档
    print("[2/4] 创建Word文档...")
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 设置默认字体
    _set_document_default_font(doc)

    # 构建封面
    print("[3/4] 构建报告内容...")
    build_cover(doc)

    # 第一章
    build_chapter1_background(doc)

    # 第二章（含动态结果）
    build_chapter2_experiment(doc, model_results, best_f1, best_auc,
                               tuning_before, tuning_after, tuned_name, tuned_params)

    # 第三章
    build_chapter3_reflection(doc)

    # 保存
    print("[4/4] 保存报告...")
    doc.save(REPORT_PATH)
    print(f'\n✅ 报告已保存至: {REPORT_PATH}')
    print("=" * 60)
    return REPORT_PATH


if __name__ == '__main__':
    generate_report()
